import os
import json
import fitz  # PyMuPDF para extração de PDFs
import requests
from bs4 import BeautifulSoup
import logging
from typing import List, Dict
from docx import Document as DocxDocument
from utils_tools.config_loader import load_config
from datetime import datetime

from uuid import uuid4
from pathlib import Path
from datetime import datetime
from openai import OpenAI
from dotenv import load_dotenv

import openai 

# from langfuse import Langfuse

def scrape_web_pages(config):
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

    print("Iniciando scraping de páginas da web...")
    urls = config["web_pages"]
    output_directory = config["storage"]["processed_texts_directory"]
    os.makedirs(output_directory, exist_ok=True)

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/119.0.0.0 Safari/537.36"
        )
    }

    for url in urls:
        try:
            try:
                response = requests.get(url, headers=headers, timeout=10, verify=True)
            except requests.exceptions.SSLError:
                print(f"⚠️ SSL falhou em {url}, tentando sem verificação...")
                response = requests.get(url, headers=headers, timeout=10, verify=False)

            if response.status_code == 200:
                soup = BeautifulSoup(response.text, "html.parser")
                text = soup.get_text(separator="\n", strip=True)

                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                slug = url.split("/")[-1] or "pagina"
                file_name = f"{slug}_{timestamp}.txt"
                file_path = os.path.join(output_directory, file_name)

                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(text)

                logging.info(f"✅ Conteúdo da página salvo: {file_name}")
                print(f"✅ Conteúdo salvo: {file_name}")
            else:
                logging.warning(f"⚠️ Erro ao acessar {url}: Status {response.status_code}")
                print(f"⚠️ Erro ao acessar {url}: Status {response.status_code}")

        except Exception as e:
            logging.error(f"❌ Erro ao processar {url}: {e}")
            print(f"❌ Erro ao processar {url}: {e}")


def pdf_to_text(pdf_path: str) -> str:
    try:
        with fitz.open(pdf_path) as doc:
            text = ""
            for page in doc:
                text += page.get_text()
        return text
    except Exception as e:
        raise RuntimeError(f"Erro ao extrair texto de '{pdf_path}': {e}")

def docx_to_text(docx_path: str) -> str:
    try:
        doc = DocxDocument(docx_path)
        return "\\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        raise RuntimeError(f"Erro ao extrair texto de '{docx_path}': {e}")

def getclean_dir(directory: str):
    for f in os.listdir(directory):
        file_path = os.path.join(directory, f)
        if os.path.isfile(file_path):
            os.remove(file_path)
    print(f"🧹 Diretório limpo: {directory}")
    
import os
import json
from datetime import datetime
from typing import List

def process_documents(config) -> List[str]:
    base_dir = config["storage"]["raw_data"]
    output_dir = config["storage"]["processed_texts_directory"]
    do_clear = config["preprocessing"]["clean_dir"]

    os.makedirs(output_dir, exist_ok=True)
    if do_clear:
        getclean_dir(output_dir)

    supported = [".pdf", ".docx", ".json"]
    files = [f for f in os.listdir(base_dir) if os.path.splitext(f)[1].lower() in supported]

    print(f"📂 Arquivos encontrados: {files}")

    if not files:
        print(f"⚠️ Nenhum arquivo PDF, DOCX ou JSON encontrado em '{base_dir}'.")
        return []

    output_files = []

    for file_name in files:
        file_path = os.path.join(base_dir, file_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = os.path.join(output_dir, f"{os.path.splitext(file_name)[0]}_{timestamp}.txt") 

        print(f"🔍 Processando: {file_name}")

        try:
            if file_name.lower().endswith(".pdf"):
                print("📄 Detectado PDF")
                text = pdf_to_text(file_path)
            elif file_name.lower().endswith(".docx"):
                print("📄 Detectado DOCX")
                text = docx_to_text(file_path)
            elif file_name.lower().endswith(".json"):
                print("📄 Detectado JSON")
                try:
                    content = open(file_path, "r", encoding="utf-8").read()
                except UnicodeDecodeError:
                    content = open(file_path, "r", encoding="utf-8-sig").read()

                data = json.loads(content)

                if isinstance(data, list):
                    text = "\n\n".join([
                        f"ID: {item.get('id')}\nCategoria: {item.get('categoria')}\nContexto: {item.get('contexto')}\nPergunta: {item.get('pergunta')}\nResposta: {item.get('resposta')}\nNext: {item.get('next')}"
                        for item in data if isinstance(item, dict)
                    ])
                else:
                    text = json.dumps(data, indent=2, ensure_ascii=False)
            else:
                print(f"⚠️ Formato não suportado: {file_name}")
                continue

            with open(output_file, "w", encoding="utf-8") as f:
                f.write(text)

            output_files.append(output_file)
            print(f"✅ Gerado TXT: {output_file}")

        except Exception as e:
            print(f"❌ Erro ao processar {file_name}: {e}")

    return output_files


def extract_text_from_file(filepath):
    ext = filepath.suffix.lower()
    if ext == ".pdf":
        return pdf_to_text(str(filepath))
    elif ext == ".txt":
        return filepath.read_text(encoding="utf-8")
    elif ext == ".docx":
        doc = DocxDocument(filepath)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif ext == ".json":
        data = json.loads(filepath.read_text(encoding="utf-8"))
        return "\n".join([str(v) for d in data for v in d.values() if isinstance(v, str)])
    else:
        return ""

def validate_output_schema(data):
    required_fields = {"id", "categoria", "contexto", "pergunta", "resposta", "next"}
    if not isinstance(data, list):
        raise ValueError("A estrutura de FAQ deve ser uma lista.")
    for i, item in enumerate(data):
        if not isinstance(item, dict):
            raise ValueError(f"Entrada {i} não é um dicionário.")
        if not required_fields.issubset(item):
            raise ValueError(f"Entrada {i} está incompleta. Esperado: {required_fields}, encontrado: {item.keys()}")


def gerar_faq_vetorizada_gpt(config):
    load_dotenv()
    raw_data_path = Path(config["storage"]["raw_data"])
    output_base = Path(config["faq_documents"][0])
    output_path = output_base / "faq_vetorizada_gpt.json"
    gpt_cfg = config["gpt"]
    output_base.mkdir(parents=True, exist_ok=True)

    openai.api_key = os.getenv("OPENAI_API_KEY")

    # langfuse = Langfuse(
    #     secret_key=os.getenv("LANGFUSE_SECRET_KEY"),
    #     public_key=os.getenv("LANGFUSE_PUBLIC_KEY"),
    #     host=os.getenv("LANGFUSE_HOST")
    # )

    system_prompt = (
        """Sua tarefa é transformar um conjunto de perguntas e respostas extraídas de documentos institucionais da Insightiva
        em entradas JSON com os campos: id, categoria, contexto, pergunta, resposta e next.
        Use linguagem simples, objetiva e focada em instruir o leitor.
        Classifique cada pergunta em uma categoria relevante e indique o próximo ID se for sequencial.
        Exemplo:
        [{
            "id": "75268040-933a-4c52-a44b-0b7d982b770d",
            "categoria": "Diagnóstico de Cultura Organizacional",
            "contexto": "Processo completo de aplicação do diagnóstico cultural, incluindo cadastro, onboarding, coleta de dados e entrega de resultados.",
            "pergunta": "Como funciona o Diagnóstico de Cultura Organizacional da Insightiva?",
            "resposta": "O diagnóstico avalia valores, comportamentos e práticas na empresa, gerando relatórios que indicam pontos fortes e oportunidades de evolução cultural.",
            "next": "37c92dcc-53d0-4cc3-bf2a-e6c765610d1f"
        }]"""
    )

    faq_items = []
    # trace = langfuse.trace(name="gerar_faq_vetorizada_gpt", user_id="sistema")

    for file in raw_data_path.iterdir():
        if not file.suffix.lower() in [".pdf", ".txt", ".docx", ".json"]:
            continue

        # span_file = trace.span(name=f"processar_arquivo_{file.name}")
        print(f"📄 Processando: {file.name}")

        try:
            texto = extract_text_from_file(file)
            chunks = [p for p in texto.split("\n") if 50 < len(p) < 300]

            for i, trecho in enumerate(chunks):
                q_id = str(uuid4())
                next_id = str(uuid4()) if i + 1 < len(chunks) else None

                # span_trecho = trace.span(name=f"gerar_faq_trecho_{i}", metadata={"arquivo": file.name})

                try:
                    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

                    response = client.chat.completions.create(
                        model=gpt_cfg["type"],
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": trecho}
                        ],
                        temperature=gpt_cfg.get("temperature", 0.3),
                        top_p=gpt_cfg.get("top_p", 0.7),
                        max_tokens=gpt_cfg.get("max_tokens", 400),
                    )

                    resposta = response.choices[0].message.content.strip()

                    faq_items.append({
                        "id": q_id,
                        "categoria": "Geral",
                        "contexto": "FAQ gerado automaticamente com base em documentos institucionais da Insightiva.",
                        "pergunta": trecho,
                        "resposta": resposta,
                        "next": next_id if next_id else None
                    })

                    # span_trecho.end(output="ok")

                except Exception as e_trecho:
                    # span_trecho.end(output=str(e_trecho), level="ERROR")
                    print(f"❌ Erro ao gerar FAQ do trecho {i}: {e_trecho}")

            # span_file.end(output="ok")

        except Exception as e_arquivo:
            # span_file.end(output=str(e_arquivo), level="ERROR")
            print(f"❌ Erro ao processar {file.name}: {e_arquivo}")

    # trace.end()

    try:
        validate_output_schema(faq_items)
        print("✅ Estrutura validada com sucesso.")
    except ValueError as ve:
        print(f"❌ Erro de validação: {ve}")
        raise

    # with open(output_path, "w", encoding="utf-8") as f:
    #     json.dump(faq_items, f, ensure_ascii=False, indent=2)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = output_base / f"faq_vetorizada_gpt_{timestamp}.json"
    with open(backup_path, "w", encoding="utf-8") as f:
        json.dump(faq_items, f, ensure_ascii=False, indent=2)

    print(f"✅ Arquivos gerados: {output_path} e {backup_path}")
    return faq_items



def gerar_e_salvar_faqs_gpt(config, limpar_diretorio=False):
    output_path = Path(config["storage"]["processed_texts_directory"])
    output_path.mkdir(parents=True, exist_ok=True)

    # 🧹 Limpa o diretório se solicitado
    if limpar_diretorio:
        getclean_dir(output_path)

    # 1. Gera a FAQ usando GPT
    faq_items = gerar_faq_vetorizada_gpt(config)

    # 2. Valida a estrutura
    try:
        validate_output_schema(faq_items)
        print("✅ Estrutura validada.")
    except ValueError as e:
        print(f"❌ Erro na estrutura dos dados: {e}")
        return

    # 3. Salva o arquivo JSON no diretório final
    nome_saida = f"faq_gerada_gpt_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    caminho_saida = output_path / nome_saida

    with open(caminho_saida, "w", encoding="utf-8") as f:
        json.dump(faq_items, f, ensure_ascii=False, indent=2)

    print(f"✅ FAQ gerada e salva em: {caminho_saida}")

